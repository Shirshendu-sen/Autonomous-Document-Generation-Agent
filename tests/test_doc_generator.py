"""
tests/test_doc_generator.py
-----------------------------
Covers Step 9: rendering the plan + drafted sections into a real .docx --
cover page, heading text, the TOC field's raw OOXML instruction, table
rows with header shading, the assumptions box, and the "refined by the
agent's self-check pass" note on revised sections.
"""
import zipfile

import pytest
from docx import Document

from app import config
from app.doc_generator import build_docx
from app.schemas import Plan, PlanSection, SectionContent


@pytest.fixture(autouse=True)
def _redirect_output_dir(tmp_path, monkeypatch):
    monkeypatch.setattr(config, "OUTPUT_DIR", tmp_path)


def _sample_plan() -> Plan:
    return Plan(
        document_type="project_plan",
        title="Mobile App Launch",
        audience="Leadership",
        assumptions=["No budget was specified.", "No firm deadline was given."],
        sections=[
            PlanSection(id="overview", title="Overview", goal="Summarise the project."),
            PlanSection(id="risks", title="Risks", goal="List risks.",
                        table_columns=["Risk", "Impact", "Mitigation"]),
        ],
    )


def _sample_sections():
    return [
        SectionContent(id="overview", title="Overview",
                        content="Some prose about the project.\nA second paragraph.",
                        revised=True),
        SectionContent(id="risks", title="Risks",
                        content=[{"Risk": "Delay", "Impact": "High", "Mitigation": "Add buffer time"}],
                        table_columns=["Risk", "Impact", "Mitigation"]),
    ]


def test_build_docx_creates_a_file(tmp_path):
    path = build_docx(_sample_plan(), _sample_sections(), "abcd1234-0000-0000-0000-000000000000")
    assert path.exists()
    assert path.suffix == ".docx"
    assert path.parent == tmp_path


def test_build_docx_contains_title_and_section_headings():
    path = build_docx(_sample_plan(), _sample_sections(), "req-1")
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    assert any("Mobile App Launch" in t for t in texts)
    assert "Overview" in texts
    assert "Risks" in texts
    assert any("Prepared for: Leadership" in t for t in texts)


def test_build_docx_contains_toc_field_instruction():
    path = build_docx(_sample_plan(), _sample_sections(), "req-2")
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert "TOC" in xml
    assert "Right-click here" in xml


def test_build_docx_contains_page_number_field():
    path = build_docx(_sample_plan(), _sample_sections(), "req-3")
    with zipfile.ZipFile(path) as z:
        footer_files = [n for n in z.namelist() if n.startswith("word/footer")]
        assert footer_files
        footer_xml = z.read(footer_files[0]).decode("utf-8")
    assert "PAGE" in footer_xml


def test_build_docx_renders_table_rows_correctly():
    path = build_docx(_sample_plan(), _sample_sections(), "req-4")
    doc = Document(path)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    header_texts = [c.text for c in table.rows[0].cells]
    assert header_texts == ["Risk", "Impact", "Mitigation"]
    assert len(table.rows) == 2  # header + 1 data row
    data_texts = [c.text for c in table.rows[1].cells]
    assert data_texts == ["Delay", "High", "Add buffer time"]


def test_build_docx_marks_revised_sections():
    path = build_docx(_sample_plan(), _sample_sections(), "req-5")
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    assert any("refined by the agent's self-check pass" in t for t in texts)


def test_build_docx_includes_assumptions_box():
    path = build_docx(_sample_plan(), _sample_sections(), "req-6")
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    assert "Assumptions Made by the Agent" in texts
    assert "No budget was specified." in texts
    assert "No firm deadline was given." in texts


def test_build_docx_omits_assumptions_box_when_there_are_none():
    plan = _sample_plan()
    plan.assumptions = []
    path = build_docx(plan, _sample_sections(), "req-7")
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    assert "Assumptions Made by the Agent" not in texts


def test_build_docx_splits_multi_paragraph_prose_into_separate_paragraphs():
    path = build_docx(_sample_plan(), _sample_sections(), "req-8")
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    assert "Some prose about the project." in texts
    assert "A second paragraph." in texts
